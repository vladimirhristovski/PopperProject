import os
import re
import time
import socket
import threading
import json
from http.server import HTTPServer, BaseHTTPRequestHandler
from socketserver import ThreadingMixIn


class ThreadedHTTPServer(ThreadingMixIn, HTTPServer):
    daemon_threads = True


from datetime import datetime
from popper import Popper
from config import (
    LOCAL_MODEL, LOCAL_PORT, LOCAL_HOST,
    ALPHA, MAX_TESTS, TIME_LIMIT, MAX_RETRY,
    DATA_DIR, HYPOTHESES, HF_TOKEN
)

os.environ["HF_TOKEN"] = HF_TOKEN
os.environ["HUGGING_FACE_HUB_TOKEN"] = HF_TOKEN

_model = None
_tokenizer = None


def load_model():
    global _model, _tokenizer
    print(f"Loading model: {LOCAL_MODEL}")
    print("This may take several minutes on first run...")

    from transformers import AutoTokenizer, AutoModelForCausalLM
    import torch

    _tokenizer = AutoTokenizer.from_pretrained(LOCAL_MODEL)
    _model = AutoModelForCausalLM.from_pretrained(
        LOCAL_MODEL,
        torch_dtype=torch.bfloat16,
        device_map="auto",
    )
    print("Model loaded successfully!")


class OpenAIHandler(BaseHTTPRequestHandler):

    def log_message(self, format, *args):
        pass

    def do_GET(self):
        if self.path == "/v1/models":
            body = json.dumps({
                "object": "list",
                "data": [{"id": LOCAL_MODEL, "object": "model"}]
            }).encode()
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", len(body))
            self.end_headers()
            self.wfile.write(body)
        else:
            self.send_response(404)
            self.end_headers()

    def do_POST(self):
        if self.path == "/v1/chat/completions":
            length = int(self.headers.get("Content-Length", 0))
            raw = self.rfile.read(length)
            data = json.loads(raw)

            messages = data.get("messages", [])
            max_tokens = data.get("max_tokens", 1024)
            temperature = data.get("temperature", 0.7)

            import torch

            text = _tokenizer.apply_chat_template(
                messages,
                tokenize=False,
                add_generation_prompt=True
            )
            inputs = _tokenizer(text, return_tensors="pt").to(_model.device)

            with torch.no_grad():
                outputs = _model.generate(
                    **inputs,
                    max_new_tokens=max_tokens,
                    temperature=max(temperature, 1e-6),
                    do_sample=temperature > 0,
                    pad_token_id=_tokenizer.eos_token_id,
                )

            response_text = _tokenizer.decode(
                outputs[0][inputs["input_ids"].shape[1]:],
                skip_special_tokens=True
            )

            body = json.dumps({
                "id": "chatcmpl-direct",
                "object": "chat.completion",
                "model": LOCAL_MODEL,
                "choices": [{
                    "index": 0,
                    "message": {"role": "assistant", "content": response_text},
                    "finish_reason": "stop"
                }],
                "usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
            }).encode()

            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", len(body))
            self.end_headers()
            self.wfile.write(body)
        else:
            self.send_response(404)
            self.end_headers()


def start_direct_server():
    load_model()

    server = ThreadedHTTPServer((LOCAL_HOST, LOCAL_PORT), OpenAIHandler)
    t = threading.Thread(target=server.serve_forever, daemon=True)
    t.start()

    for _ in range(30):
        try:
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(2)
            if s.connect_ex((LOCAL_HOST, LOCAL_PORT)) == 0:
                s.close()
                print("Server is ready!")
                return True
            s.close()
        except Exception:
            pass
        time.sleep(1)

    print("ERROR: Server did not start.")
    return False


def check_data():
    required = ["winobias.csv", "bbq.csv", "stereoset.csv"]
    missing = [f for f in required if not os.path.exists(os.path.join(DATA_DIR, f))]
    return missing


def parse_result(result):
    if not result or not isinstance(result, dict):
        return 0.0, "error"

    print(f"  Raw result keys: {list(result.keys())}")
    e_value = None
    decision = None

    parsed = result.get("parsed_result", {})
    if isinstance(parsed, dict):
        for key in ("e_value", "e_val", "evalue", "combined_e_value", "final_e_value", "E_value"):
            val = parsed.get(key)
            if val is not None:
                try:
                    candidate = float(val)
                    if candidate > 0:
                        e_value = candidate
                        break
                except (TypeError, ValueError):
                    pass

    last_msg = str(result.get("last_message", ""))
    log_text = str(result.get("log", ""))
    full_text = last_msg + "\n" + log_text

    if not e_value:
        m = re.search(r'e.?value[^:\n]*calibrator[^:\n]*:\s*([0-9]+(?:\.[0-9]+)?(?:[eE][+\-]?[0-9]+)?)', full_text,
                      re.IGNORECASE)
        if m:
            try:
                e_value = float(m.group(1))
            except ValueError:
                pass

    if not e_value:
        m = re.search(r'combined\s+e.?value[^:\d]*:\s*([0-9]+(?:\.[0-9]+)?(?:[eE][+\-]?[0-9]+)?)', full_text,
                      re.IGNORECASE)
        if m:
            try:
                e_value = float(m.group(1))
            except ValueError:
                pass

    if not e_value:
        m = re.search(r'\bE[-_]?[Vv]alue\s*[:\s]\s*([0-9]+(?:\.[0-9]+)?(?:[eE][+\-]?[0-9]+)?)', full_text)
        if m:
            try:
                e_value = float(m.group(1))
            except ValueError:
                pass

    text_lower = full_text.lower()
    if "sufficient evidence - pass" in text_lower:
        decision = "SUPPORTED"
    elif "insufficient evidence - continue" in text_lower:
        decision = "NOT SUPPORTED"
    else:
        for kw in ("SUPPORTED", "NOT SUPPORTED", "PASS", "CONTINUE", "REJECT", "FAIL TO REJECT"):
            if kw in full_text.upper():
                decision = kw
                break

    try:
        e_value = float(e_value) if e_value is not None else 0.0
    except (TypeError, ValueError):
        e_value = 0.0

    decision = str(decision) if decision else "unknown"
    print(f"  → Parsed  e_value={e_value:.4f}  decision={decision}")
    return e_value, decision


def determine_status(e_value, decision):
    d = decision.strip().upper()
    if "NOT SUPPORTED" in d or "CONTINUE" in d:
        return "NOT SUPPORTED"
    if "SUPPORTED" in d or "PASS" in d:
        return "SUPPORTED"
    return "SUPPORTED" if e_value >= 10 else "NOT SUPPORTED"


def generate_report(results, total_time, results_file, model_name):
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        _save_text_report(results, total_time, results_file, model_name)
        return

    report_file = results_file.replace(".csv", "_report.docx")
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
    supported = sum(1 for r in results if r["status"] == "SUPPORTED")
    not_supp = sum(1 for r in results if r["status"] == "NOT SUPPORTED")
    errors = sum(1 for r in results if r["status"] == "ERROR")

    def _cell_borders(cell):
        tc = cell._tc;
        tcPr = tc.get_or_add_tcPr()
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single");
            b.set(qn("w:sz"), "4");
            b.set(qn("w:color"), "CCCCCC")
            tcPr.append(b)

    def _header_cell(table, row, col, text):
        cell = table.cell(row, col);
        cell.text = text
        run = cell.paragraphs[0].runs[0];
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF);
        run.font.size = Pt(10)
        tc = cell._tc;
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd");
        shd.set(qn("w:fill"), "2E5090");
        tcPr.append(shd)
        _cell_borders(cell)

    def _data_cell(table, row, col, text, bg="FFFFFF"):
        cell = table.cell(row, col);
        cell.text = str(text)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if bg != "FFFFFF":
            tc = cell._tc;
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd");
            shd.set(qn("w:fill"), bg);
            tcPr.append(shd)
        _cell_borders(cell)

    STATUS_COLORS = {"SUPPORTED": "D5F5E3", "NOT SUPPORTED": "FDECEA", "ERROR": "FFF3CD"}
    CONFIG_ROWS = [("Alpha:", str(ALPHA)), ("Max tests:", str(MAX_TESTS)),
                   ("Time limit:", f"{TIME_LIMIT} min"), ("Aggregate:", "E-value"),
                   ("Datasets:", "WinoBias, BBQ, StereoSet")]

    doc = Document()
    section = doc.sections[0]
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.54))

    t = doc.add_heading("POPPER Experiment Report", 0);
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.add_run("Model: ").bold = True;
    meta.add_run(f"{model_name}    ")
    meta.add_run("Date: ").bold = True;
    meta.add_run(f"{timestamp}    ")
    meta.add_run("Total time: ").bold = True;
    meta.add_run(f"{total_time:.1f} min")

    doc.add_paragraph();
    doc.add_heading("Summary", level=1)
    s = doc.add_paragraph()
    for label, val in [("Supported: ", supported), ("Not Supported: ", not_supp),
                       ("Errors: ", errors), ("Total: ", len(results))]:
        s.add_run(label).bold = True;
        s.add_run(f"{val}   ")

    doc.add_paragraph();
    doc.add_heading("Results", level=1)
    tbl = doc.add_table(rows=len(results) + 1, cols=5);
    tbl.style = "Table Grid"
    for i, hdr in enumerate(["Hypothesis", "Status", "E-value", "Decision", "Time"]):
        _header_cell(tbl, 0, i, hdr)
    for i, r in enumerate(results, 1):
        hyp = r["hypothesis"][:80] + ("…" if len(r["hypothesis"]) > 80 else "")
        bg = STATUS_COLORS.get(r["status"], "FFFFFF")
        _data_cell(tbl, i, 0, hyp);
        _data_cell(tbl, i, 1, r["status"], bg)
        _data_cell(tbl, i, 2, f"{r['e_value']:.4f}" if r["e_value"] else "N/A")
        _data_cell(tbl, i, 3, str(r["decision"])[:60]);
        _data_cell(tbl, i, 4, f"{r['time_min']:.1f} min")

    doc.add_paragraph();
    doc.add_heading("Configuration", level=1)
    cfg_tbl = doc.add_table(rows=len(CONFIG_ROWS), cols=2);
    cfg_tbl.style = "Table Grid"
    for idx, (key, val) in enumerate(CONFIG_ROWS):
        k = cfg_tbl.cell(idx, 0);
        k.text = key;
        k.paragraphs[0].runs[0].bold = True;
        _cell_borders(k)
        v = cfg_tbl.cell(idx, 1);
        v.text = val;
        _cell_borders(v)

    doc.add_paragraph();
    doc.add_heading("Individual Test Details", level=1)
    for i, r in enumerate(results, 1):
        doc.add_heading(f"Test {i}", level=2)
        for label, val in [("Hypothesis", r["hypothesis"]), ("Status", r["status"]),
                           ("E-value", f"{r['e_value']:.6f}"), ("Decision", r["decision"]),
                           ("Time", f"{r['time_min']:.1f} min")]:
            doc.add_paragraph(f"{label:12}: {val}")
        if i < len(results): doc.add_paragraph("─" * 50)

    doc.add_page_break();
    doc.add_heading("Final Summary", level=1)
    avg_e = sum(r["e_value"] for r in results) / len(results) if results else 0
    doc.add_paragraph().add_run(
        f"{'SUPPORTED' if supported > 0 else 'NOT SUPPORTED':15} | avg E={avg_e:.4f} | {total_time:.1f}m | {model_name}")
    doc.add_paragraph()
    for label, val in [("Total time : ", f"{total_time:.1f} min"), ("Supported  : ", str(supported)),
                       ("Not Supp.  : ", str(not_supp)), ("Errors     : ", str(errors))]:
        p = doc.add_paragraph();
        p.add_run(label).bold = True;
        p.add_run(val)

    doc.save(report_file)
    print(f"Word report saved to: {os.path.abspath(report_file)}")


def _save_text_report(results, total_time, results_file, model_name):
    txt = results_file.replace(".csv", "_report.txt")
    supported = sum(1 for r in results if r["status"] == "SUPPORTED")
    not_supp = sum(1 for r in results if r["status"] == "NOT SUPPORTED")
    errors = sum(1 for r in results if r["status"] == "ERROR")
    with open(txt, "w") as f:
        f.write(f"POPPER Experiment Report — {model_name}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total time: {total_time:.1f} min\n")
        f.write("=" * 70 + "\n")
        f.write(f"Supported: {supported}  Not Supported: {not_supp}  Errors: {errors}  Total: {len(results)}\n\n")
        for r in results:
            f.write(f"Hypothesis : {r['hypothesis']}\n")
            f.write(f"Status     : {r['status']}\n")
            f.write(f"E-value    : {r['e_value']:.6f}\n")
            f.write(f"Decision   : {r['decision']}\n")
            f.write(f"Time       : {r['time_min']:.1f} min\n")
            f.write("-" * 50 + "\n\n")
    print(f"Text report saved to: {txt}")


def run(results_file="results_direct.csv"):
    print("=" * 70)
    print(f"POPPER — Test 3 Direct: ({LOCAL_MODEL})")
    print("=" * 70)

    missing = check_data()
    if missing:
        print(f"ERROR: Missing data files: {missing}")
        print("Run: python3 prepare_data.py")
        return

    if not start_direct_server():
        print("ERROR: Could not start server.")
        return

    os.environ["OPENAI_API_KEY"] = "direct"
    os.environ["OPENAI_BASE_URL"] = f"http://{LOCAL_HOST}:{LOCAL_PORT}/v1"

    results = []
    total_start = time.time()

    print(f"\nRunning {len(HYPOTHESES)} hypotheses\n")
    print("=" * 70)

    for i, hypothesis in enumerate(HYPOTHESES, 1):
        print(f"\n[{i}/{len(HYPOTHESES)}] {hypothesis[:65]}...")
        print("  Initializing POPPER agent...")

        try:
            agent = Popper(llm=LOCAL_MODEL, is_locally_served=True, server_port=LOCAL_PORT)
            agent.register_data(data_path=DATA_DIR, loader_type="custom")
            agent.configure(
                alpha=ALPHA,
                max_num_of_tests=MAX_TESTS,
                max_retry=MAX_RETRY,
                time_limit=TIME_LIMIT,
                aggregate_test="E-value",
                relevance_checker=True,
                use_react_agent=True,
            )
        except Exception as e:
            print(f"  ERROR: Failed to initialize agent — {e}")
            results.append({"model": LOCAL_MODEL, "hypothesis": hypothesis,
                            "status": "ERROR", "e_value": 0.0,
                            "decision": "init_error", "time_min": 0.0})
            continue

        start = time.time()
        try:
            result = agent.validate(hypothesis=hypothesis)
            elapsed = (time.time() - start) / 60
            e_value, decision = parse_result(result)
            status = determine_status(e_value, decision)

            print(f"  Status  : {status}")
            print(f"  E-value : {e_value:.4f}")
            print(f"  Decision: {decision}")
            print(f"  Time    : {elapsed:.1f} min")

            results.append({"model": LOCAL_MODEL, "hypothesis": hypothesis,
                            "status": status, "e_value": e_value,
                            "decision": decision, "time_min": elapsed})

        except Exception as e:
            elapsed = (time.time() - start) / 60
            print(f"  ERROR: {str(e)[:120]}")
            results.append({"model": LOCAL_MODEL, "hypothesis": hypothesis,
                            "status": "ERROR", "e_value": 0.0,
                            "decision": "error", "time_min": elapsed})

    total_time = (time.time() - total_start) / 60

    import pandas as pd
    df = pd.DataFrame(results)
    df.to_csv(results_file, index=False)
    print(f"\nCSV saved to: {results_file}")

    generate_report(results, total_time, results_file, LOCAL_MODEL)

    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    for r in results:
        print(f"{r['status']:15} | E={r['e_value']:7.4f} | {r['time_min']:5.1f}m | {r['hypothesis'][:40]}")
    print("=" * 70)
    print(f"Total time : {total_time:.1f} min ({total_time / 60:.2f} hours)")
    print(f"Supported  : {sum(1 for r in results if r['status'] == 'SUPPORTED')}")
    print(f"Not Supp.  : {sum(1 for r in results if r['status'] == 'NOT SUPPORTED')}")
    print(f"Errors     : {sum(1 for r in results if r['status'] == 'ERROR')}")
    print("=" * 70)


if __name__ == "__main__":
    run()
